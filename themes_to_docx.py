import json
import os
import re
from docx import Document
import pandas as pd
from tqdm import tqdm


def read_config(json_file_path=None):
    with open(json_file_path, 'r', encoding="utf-8") as file:
        data = json.load(file)
    return data


def extract_substring(s):
    pos = s.rfind(", ")
    if pos == -1:
        return s  # Если ', ' не найдено, возвращаем всю строку
    else:
        return s[pos+2:]  # Возвращаем подстроку после последнего ', '


def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)
    return 0


def replace_n_save(
    input_doc="Титульный лист КР.docx",
    SurName="Иванов",
    FirstName="Иван",
    SecondName="Иванович",
    Group="ИИИИ-01-01",
    Theme="Анализ рынка",
    MainTeacher="Юрч",
    out_doc="ИИИИ-01-01_ИвановИИ.docx"):

    regex_list = (
        re.compile(r"Тема\ работы"),
        re.compile(r"Иванов\ Иван\ Иванович"),
        re.compile(r"ИИИИ-01-01"),
        re.compile(r"Руководитель\ курсовой\ работы"),
        re.compile(r"Иванов\ И.И."),
        re.compile(r"Петров\ П.П.")
        )

    doc = Document(input_doc)
    docx_replace_regex(doc, regex_list[0], Theme)
    docx_replace_regex(doc, regex_list[1], rf"{SurName} {FirstName} {SecondName}")
    docx_replace_regex(doc, regex_list[2], Group)
    docx_replace_regex(doc, regex_list[3], MainTeacher)
    docx_replace_regex(doc, regex_list[4], rf"{SurName} {FirstName[0]}.{SecondName[0]}.")
    docx_replace_regex(doc, regex_list[5], extract_substring(MainTeacher))
    doc.save(out_doc)
    return 0


def main_routine():
    # Читаем пути файлов
    config = read_config("config.json")

    """
    Устанавливаем перечень названий полей читаемого табличного файла
    
    SurName - Фамилия студента
    FirstName - Имя студента
    SecondName - Отчество студента
    Group - Название группы
    Theme - тема работы
    MainTeacher - степень, звание, должность, Фамилия И.О. (руководителя)
    """
    columns_df = ["SurName", "FirstName", "SecondName", "Group", "Theme", "MainTeacher"]
    df = pd.read_excel(config.get("from_xlsx"), names=columns_df)

    # Забираем откуда читаем шаблон и куда кладем заполненные документы
    input_docx = config.get("from_docx")
    out_dir = config.get("to_dir")

    # Итерируемся по всем строкам таблицы с темами студентов
    for idx, row in tqdm(enumerate(df.to_dict(orient='records'))):
        # Формат названия файла: ИИИИ-01-01_Фамилия_Задание.docx
        out_filename = f"{row.get('Group')}_" + \
                       f"{row.get('SurName')}" + \
                       f"{row.get('FirstName')[0]}{row.get('SecondName')[0]}_" +\
                       f"{config.get('last_in_filename')}.docx"

        out_docx = os.path.join(out_dir, out_filename)
        replace_n_save(input_doc=input_docx,
                       **row,
                       out_doc=out_docx)
    return 0


if __name__ == "__main__":
    main_routine()